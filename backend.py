import os
from pdf2docx import Converter
from docx import Document
import re
from PyPDF2 import PdfReader
import xlsxwriter
import aspose.words as aw
def identify_filetypes(folder_path):
    pdf_files = []
    docx_files = []
    doc_files = []
    # Iterate over files in the folder
    for filename in os.listdir(folder_path):
        filepath = os.path.join(folder_path, filename)
        # Check if it's a file (not a directory)
        if os.path.isfile(filepath):
            # Check file extension
            if filename.lower().endswith('.pdf'):
                pdf_files.append(filename)
            elif filename.lower().endswith('.docx'):
                docx_files.append(filename)
            elif filename.lower().endswith('.doc'):
                doc_files.append(filename)

    return [pdf_files,docx_files,doc_files]

def doc_to_docx(files,content_path):
    for doc_file in files:
        file_name=os.path.splitext(doc_file)[0]
        docx_name=file_name + ".docx"

        doc_path = os.path.join(content_path, doc_file)
        docx_path = os.path.join(content_path, docx_name)

        doc = aw.Document(doc_path)
        doc.save(docx_path)
        os.remove(doc_path)

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def extract_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, "rb") as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text


def extract_text_between_words(text, word1, word2):
    pattern = re.compile(r'{}(.*?){}'.format(re.escape(word1), re.escape(word2)),re.DOTALL)
    match = pattern.search(text)
    if match:
        return match.group(1).strip()
    else:
        return None


def extract_emails_old(text):
    # Define the regular expression pattern for email addresses
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

    # Find all email addresses in the text using the pattern
    emails = re.findall(pattern, text)

    return emails

def extract_emails(text):
    pattern = r'\b[\w.\d]*@\w+(?:\.\w+)*\b'
    emails = re.findall(pattern, text)
    return emails
def extract_phone_numbers(text):
    # Define the regular expression pattern for phone numbers
    pattern = r'\b(?:\+?\d{2}-)?(?:\d{3}[-\s]?\d{3}[-\s]?\d{4}|\d{5}[-\s]?\d{5})\b'

    # Find all phone numbers in the text using the pattern
    phone_numbers = re.findall(pattern, text)

    return phone_numbers
def section_extractor(text,filename):
    headings = [
        'Personal Information',
        'Objective',
        'Work History',
        'Education',
        'Work Experience',
        'Working Experience',
        'Professional Experience',
        'Skills',
        'Certifications',
        'Certificates',
        'Projects',
        'Publications',
        'Awards',
        'Employment History',
        'Professional Affiliations',
        'References',
        'Languages',
        'Achievement',
        'Academic Credentials',
        'Academic Qualification',
        'Professional Qualifications',
        'Profile',
        'Personal Details',
        'Academic Details',
        'Soft Skills',
        'Personal Skills',
        'Software Skills',
        'Strengths',
        'Tool Stack',
        'Hobbies',
        'Interests',
        'Computer Proficiency',
        'Core Competencies',
        'Internship',
        'Work Summary',
        'Desired Job Details',
        'Professional Interaction',
        'Professional Summary',
        'Summary',
        'Educational Details',
        'Details'


    ]

    text_data = text.split()
    sections_present=[]

    "---------Removing Aspose watermark-----------------"
    if text_data[4] == "Aspose.Words.":
        for i in range(10):
            text_data.pop(0)
    "---------------"

    sections_present.append(text_data[0])
    for word in headings:
        capital=word.upper()
        if capital in text:
            sections_present.append(capital)
        elif word in text:
            sections_present.append(word)


    "--------------------------------------------"
    "special case handling"

    if 'Work Experience' not in sections_present and 'Work Experience'.upper() not in sections_present and \
            'Professional Experience' not in sections_present and 'Professional Experience'.upper() not in sections_present:
        if "EXPERIENCE" in text:
            sections_present.append("EXPERIENCE")
        elif "Experience" in text:
            sections_present.append("Experience")

    if "ACADEMIC CREDENTIALS" in sections_present and "Education" in sections_present:
        sections_present.remove("Education") # special case for Akash Goel

    "------------------------------------------------------------------------------------------------------"
    extracted_data=[]

    base_name, extension = os.path.splitext(filename)

    phone_numbers=extract_phone_numbers(text=text)
    phone_numbers_str=" ,".join(phone_numbers)

    emails=extract_emails(text=text)
    emails_str=" ,".join(emails)

    extracted_data.append(["Name",base_name])
    extracted_data.append(["Contact Number",phone_numbers_str])
    extracted_data.append(["Email",emails_str])

    for section in sections_present:

        min_data=None

        for next_section in sections_present:
            if section != next_section:
                data = extract_text_between_words(text=text, word1=f"{section}", word2=f"{next_section}")

                if min_data:
                    if data:
                        if len(data) < len(min_data) and len(data) > 0:
                            min_data=data
                else:
                    if data:
                        min_data=data

        if not min_data:
            "now consider its the last section thus we wil match till end of doc"
            end_pattern = re.compile(r'{}(.*)'.format(re.escape(section)), re.DOTALL)
            match = end_pattern.search(text)
            if match:
                min_data = match.group(1).strip()

        extracted_data.append([section,min_data])

    """for item in extracted_data:
        print(item[0])
        print(item[1])"""
    return extracted_data



def excel_writer(data):
    if os.path.exists("Report.xlsx"):
        os.remove("Report.xlsx")
    workbook = xlsxwriter.Workbook(f"Report.xlsx")



    for CV in data:
        worksheet = workbook.add_worksheet()
        col = 0
        for index,item in enumerate(CV):
            if index == 3:

                worksheet.write(0, col, f"Introduction")
                worksheet.write(1, col, f"{item[0]} {item[1]}")
                col += 1
            else:
                worksheet.write(0, col, f"{item[0]} ")
                worksheet.write(1, col, f"{item[1]}")
                col += 1

        worksheet.write(3,0,"Kindly")
        worksheet.write(3,1,"double")
        worksheet.write(3,2,"click")
        worksheet.write(3,3,"on")
        worksheet.write(3,4,"the")
        worksheet.write(3,5,"cells")
        worksheet.write(3,6,"to")
        worksheet.write(3,7,"view")
        worksheet.write(3,8,"its")
        worksheet.write(3,9,"entire")
        worksheet.write(3,10,"content")





    workbook.close()
    return 1


def clear_directory(directory):
    # Iterate over all the files and subdirectories in the specified directory
    for file_or_dir in os.listdir(directory):
        # Construct the full path of the file or subdirectory
        full_path = os.path.join(directory, file_or_dir)

        # Check if it's a file
        if os.path.isfile(full_path):
            # If it's a file, remove it
            os.remove(full_path)
        elif os.path.isdir(full_path):
            # If it's a directory, clear its contents recursively
            clear_directory(full_path)
            # After clearing the subdirectory, remove it
            os.rmdir(full_path)
if __name__=="__main__":
    text=extract_text_from_pdf("extracted/Sample2/CAChamanKumar.pdf")
    #text=extract_text_from_docx("extracted/Sample2/heemSen.docx")

    section_extractor(text=text)

